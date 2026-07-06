#!/usr/bin/env python3
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs" / "customer"
OUTPUT_DIR = ROOT / "output" / "documents"

FONT_LATIN = "Calibri"
FONT_CN = "PingFang SC"
INK = RGBColor(24, 38, 45)
MUTED = RGBColor(90, 104, 112)
ACCENT = RGBColor(18, 105, 122)
ACCENT_DARK = RGBColor(8, 65, 78)
WARN = RGBColor(129, 82, 12)
RED = RGBColor(141, 36, 36)

LIGHT_FILL = "F3F7F8"
HEADER_FILL = "E6F0F2"
WARN_FILL = "FFF6E6"
BLUE_FILL = "EEF6F8"
BORDER = "C9D6DA"
WHITE = "FFFFFF"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


@dataclass(frozen=True)
class DocMeta:
    title: str
    subtitle: str
    filename: str
    source_filename: str
    doc_type: str
    audience: str
    version: str
    updated: str
    lead: str


def set_run_font(
    run,
    *,
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: RGBColor | None = None,
) -> None:
    run.font.name = FONT_LATIN
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def paragraph_border_bottom(paragraph, color="2E9BB0", size="10") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)


def set_paragraph_spacing(paragraph, *, before=0, after=6, line=1.25) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_margins(cell, top=90, bottom=90, start=130, end=130) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), BORDER)

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])
            set_cell_margins(cell)


def set_cell_text(cell, text: str, *, bold=False, color=INK, size=10.4) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_paragraph_spacing(paragraph, after=0, line=1.2)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def configure_doc(doc: Document, meta: DocMeta) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(11.2)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, ACCENT_DARK, 16, 8),
        ("Heading 2", 13.3, ACCENT, 12, 6),
        ("Heading 3", 12, ACCENT_DARK, 8, 4),
    ]:
        style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.18

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run(meta.doc_type)
    set_run_font(run, size=9.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("万法常世 AI 全智能客服系统")
    set_run_font(run, size=9, color=MUTED)


def add_title_block(doc: Document, meta: DocMeta) -> None:
    kicker = doc.add_paragraph()
    set_paragraph_spacing(kicker, before=6, after=4, line=1.0)
    run = kicker.add_run(meta.doc_type)
    set_run_font(run, size=10.8, bold=True, color=ACCENT)

    title = doc.add_paragraph()
    set_paragraph_spacing(title, before=0, after=4, line=1.12)
    run = title.add_run(meta.title)
    set_run_font(run, size=24, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    set_paragraph_spacing(subtitle, after=12, line=1.18)
    run = subtitle.add_run(meta.subtitle)
    set_run_font(run, size=12.6, color=MUTED)
    paragraph_border_bottom(subtitle)

    rows = [
        ("适用对象", meta.audience),
        ("版本", meta.version),
        ("更新时间", meta.updated),
    ]
    table = doc.add_table(rows=0, cols=2)
    set_table_geometry(table, [1600, CONTENT_WIDTH_DXA - 1600])
    for label, value in rows:
        row = table.add_row()
        cell_shading(row.cells[0], HEADER_FILL)
        cell_shading(row.cells[1], WHITE)
        set_cell_text(row.cells[0], label, bold=True, color=ACCENT_DARK, size=10.2)
        set_cell_text(row.cells[1], value, color=INK, size=10.2)
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, after=2)

    add_callout(doc, "核心结论", meta.lead, fill=BLUE_FILL)


def add_para(doc: Document, text: str, *, color=INK, bold=False) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, after=7, line=1.25)
    run = paragraph.add_run(text)
    set_run_font(run, size=11.4, color=color, bold=bold)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.38)
        paragraph.paragraph_format.first_line_indent = Inches(-0.19)
        set_paragraph_spacing(paragraph, after=4, line=1.22)
        run = paragraph.add_run(item)
        set_run_font(run, size=11.1, color=INK)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.left_indent = Inches(0.38)
        paragraph.paragraph_format.first_line_indent = Inches(-0.19)
        set_paragraph_spacing(paragraph, after=5, line=1.22)
        run = paragraph.add_run(item)
        set_run_font(run, size=11.1, color=INK)


def add_heading(doc: Document, text: str, level=1) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell_shading(cell, HEADER_FILL)
        set_cell_text(cell, header, bold=True, color=ACCENT_DARK, size=10.2)
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            cell_shading(cell, LIGHT_FILL if idx == 0 else WHITE)
            set_cell_text(cell, value, bold=(idx == 0), color=INK, size=10.0)
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, after=4)


def add_callout(doc: Document, label: str, text: str, *, fill=BLUE_FILL, color=INK) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.rows[0].cells[0]
    cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=1, line=1.18)
    label_run = p.add_run(label + "：")
    set_run_font(label_run, size=10.6, bold=True, color=ACCENT_DARK if fill != WARN_FILL else WARN)
    body_run = p.add_run(text)
    set_run_font(body_run, size=10.6, color=color)
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, after=2)


def md_table(headers: list[str], rows: list[list[str]]) -> str:
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join("---" for _ in headers) + " |"]
    lines += ["| " + " | ".join(value.replace("\n", "<br>") for value in row) + " |" for row in rows]
    return "\n".join(lines)


def write_markdown(meta: DocMeta, blocks: list[tuple[str, object]]) -> Path:
    lines = [
        f"# {meta.title}",
        "",
        meta.subtitle,
        "",
        f"- 文档类型：{meta.doc_type}",
        f"- 适用对象：{meta.audience}",
        f"- 版本：{meta.version}",
        f"- 更新时间：{meta.updated}",
        "",
        f"> 核心结论：{meta.lead}",
        "",
    ]
    for kind, payload in blocks:
        if kind == "h1":
            lines += [f"## {payload}", ""]
        elif kind == "h2":
            lines += [f"### {payload}", ""]
        elif kind == "p":
            lines += [str(payload), ""]
        elif kind == "bullets":
            lines += [f"- {item}" for item in payload] + [""]
        elif kind == "numbered":
            lines += [f"{idx}. {item}" for idx, item in enumerate(payload, 1)] + [""]
        elif kind == "table":
            headers, rows = payload
            lines += [md_table(headers, rows), ""]
        elif kind == "callout":
            label, text = payload
            lines += [f"> **{label}**：{text}", ""]
        elif kind == "page_break":
            lines += ["---", ""]
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    path = DOCS_DIR / meta.source_filename
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def build_docx(meta: DocMeta, blocks: list[tuple[str, object]]) -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_doc(doc, meta)
    add_title_block(doc, meta)
    for kind, payload in blocks:
        if kind == "h1":
            add_heading(doc, str(payload), 1)
        elif kind == "h2":
            add_heading(doc, str(payload), 2)
        elif kind == "p":
            add_para(doc, str(payload))
        elif kind == "bullets":
            add_bullets(doc, payload)
        elif kind == "numbered":
            add_numbered(doc, payload)
        elif kind == "table":
            headers, rows, widths = payload
            add_table(doc, headers, rows, widths)
        elif kind == "callout":
            label, text, fill = payload
            add_callout(doc, label, text, fill=fill)
        elif kind == "page_break":
            doc.add_page_break()
    path = OUTPUT_DIR / meta.filename
    doc.save(path)
    return path


def ops_blocks() -> list[tuple[str, object]]:
    return [
        ("h1", "一、两种正式运营模式"),
        (
            "p",
            "正式交付后，智能客服系统不应再依赖本地隧道、临时端口或人工口头配置。运营模式固定收口为两条线：由我方统一运行的托管云端版，以及部署在客户环境中的私有化部署版。两种模式都需要账号体系、权限、审计、备份、知识库运营、模型成本控制和故障响应流程。",
        ),
        (
            "table",
            (
                ["模式", "适合对象", "客户怎么使用", "我方怎么维护"],
                [
                    ["托管云端版", "中小企业、首批试点客户、希望快速上线的客户", "客户访问我方提供的正式域名，使用账号登录客服中台", "我方负责云服务器、数据库、备份、升级、监控、排障和模型/渠道配置"],
                    ["私有化部署版", "有内控、合规、数据本地化或内网要求的客户", "客户访问自己的内网域名、专有云域名或客户公网域名", "我方交付部署包和运维 SOP，通过诊断包、临时授权、变更单和权限回收进行维护"],
                ],
                [1700, 2500, 2450, 2710],
            ),
        ),
        ("callout", ("原则", "托管云端版卖的是持续运行能力；私有化部署版卖的是系统交付、客户环境落地和受控维护能力。两种模式都不能把本地测试链路包装成正式上线。", BLUE_FILL)),
        ("h1", "二、正式商用必须具备的账号体系"),
        ("p", "客户购买系统后，实际进入的是客服中台，而不是直接进入企业微信后台。企业微信、公众号、抖音、电商平台只是消息入口；我们的客服中台负责会话、知识库、AI 草稿、人工审核、质量复盘和运营管理。"),
        (
            "table",
            (
                ["账号能力", "最低要求", "建议做法"],
                [
                    ["企业空间", "每个客户至少一个独立租户", "托管云端版先采用单客户单实例；后续多租户 SaaS 再单独立项"],
                    ["登录身份", "邮箱和密码登录", "邮箱用于邀请、找回密码、通知和审计；后续可接企业微信或飞书登录"],
                    ["角色权限", "管理员、运营负责人、坐席、审核员、只读审计", "角色决定能否管理知识库、审核草稿、查看报表、改配置和导出诊断"],
                    ["成员管理", "邀请、禁用、重置密码", "离职成员必须可停用；客户管理员可管理坐席和运营成员"],
                    ["审计日志", "登录、配置、知识库、审核、外发计划都要留痕", "后续合同可约定日志保留周期和导出方式"],
                    ["安全增强", "强密码、会话过期、敏感配置不展示明文", "企业增强版再加 MFA、单点登录、IP 白名单和操作审批"],
                ],
                [1600, 2700, 5060],
            ),
        ),
        ("h1", "三、托管云端版运营模式"),
        ("h2", "1. 部署形态"),
        ("bullets", [
            "近期建议采用“单客户单实例”：一个客户对应一套独立应用、数据库、Redis 和配置。",
            "正式域名由我方或客户提供，必须使用 HTTPS。",
            "平台回调地址使用正式域名，不使用 localtunnel、ngrok 这类临时隧道。",
            "模型 Key、渠道 Token、EncodingAESKey 等密钥进入云端环境变量或 Secret 管理，不写入文档和聊天。",
            "外发默认关闭，进入真实发送前必须有白名单、人工审核、发送门禁和客户授权。",
        ]),
        ("h2", "2. 运营职责"),
        (
            "table",
            (
                ["事项", "我方负责", "客户负责"],
                [
                    ["基础设施", "服务器、容器、数据库、Redis、备份、监控、日志", "提供域名或确认使用我方子域名"],
                    ["账号体系", "创建客户空间、初始管理员、角色模板", "指定管理员、坐席、知识负责人"],
                    ["知识库", "提供导入模板、检索配置、质量复盘工具", "提供产品资料、价格政策、售后规则和禁用话术"],
                    ["模型", "配置模型路由、限额、fallback、成本统计", "确认模型费用归属和调用范围"],
                    ["渠道", "配置回调、验签、入站、草稿和 outbox", "提供官方账号权限、测试账号和授权材料"],
                    ["故障", "监控、诊断、修复、升级、复盘", "配合确认业务影响和测试结果"],
                ],
                [1800, 3900, 3660],
            ),
        ),
        ("h2", "3. 日常运维节奏"),
        ("numbered", [
            "每日看健康状态、错误日志、模型调用失败、渠道回调失败和队列积压。",
            "每周看低置信问题、人工改写问题、知识缺口和高风险话术。",
            "每月输出运营复盘，包括接待量、AI 草稿采纳率、转人工率、知识库更新、典型失败问题和下月优化计划。",
            "每次版本升级前先备份，再在低峰窗口更新，升级后跑健康检查、登录、知识检索、AI 草稿、人审和 outbox 门禁。",
        ]),
        ("h1", "四、私有化部署版运营模式"),
        ("h2", "1. 交付形态"),
        ("p", "私有化不是把源码发给客户，也不是把当前开发目录复制到客户服务器。正式交付应是版本化部署包，包含镜像、编排文件、环境模板、迁移说明、备份恢复、诊断脚本、客户手册、验收清单和权限回收记录。"),
        (
            "table",
            (
                ["交付物", "内容", "验收要点"],
                [
                    ["镜像包", "后端、前端、数据库/Redis 依赖镜像或客户环境适配说明", "镜像 tag、校验值、版本号一致"],
                    ["部署编排", "Docker Compose 或 Kubernetes 清单", "服务可启动，健康检查可用，数据库和 Redis 不直接暴露公网"],
                    ["环境模板", "变量名、说明和安全默认值", "不包含真实 API Key、Token、Secret 或客户密码"],
                    ["迁移说明", "Alembic revision、升级窗口、备份和回滚方式", "非空库迁移必须有备份和客户授权"],
                    ["诊断工具", "只读诊断包生成脚本", "不输出 .env 明文、密钥、客户聊天原文"],
                    ["验收表", "健康、登录、知识检索、AI 草稿、人审、outbox 门禁", "客户确认后才算交付完成"],
                ],
                [1600, 4100, 3660],
            ),
        ),
        ("h2", "2. 私有化维护流程"),
        ("numbered", [
            "客户先运行脱敏诊断包，优先离线分析。",
            "诊断包不足时，客户开临时只读远程通道，例如 VPN、零信任、堡垒机或屏幕共享。",
            "我方只执行本次故障需要的只读检查，不读取 .env 明文，不导出客户原始数据。",
            "如需重启、迁移、改配置、恢复数据、打开外发，必须提交二次授权。",
            "变更前确认备份和回滚点。",
            "变更后跑健康检查、登录、知识检索、AI 草稿、人审和 outbox 门禁。",
            "排障结束后客户回收账号、关闭隧道、撤销临时密钥，必要时轮换凭据。",
            "P0/P1 故障必须形成复盘，记录影响、时间线、原因、修复、验证和预防措施。",
        ]),
        ("callout", ("安全边界", "私有化维护不是长期后门。我方不长期保存客户服务器密码、数据库密码、模型 Key、平台 Token 或 webhook secret。", WARN_FILL)),
        ("h1", "五、两种模式的费用和成本口径"),
        (
            "table",
            (
                ["成本项", "托管云端版", "私有化部署版"],
                [
                    ["服务器和数据库", "通常由我方纳入月费或单独实报实销", "通常由客户承担；我方只负责部署和维护指导"],
                    ["模型调用", "可由我方代管并计入额度，也可客户 BYOK", "优先客户 BYOK 或客户专属模型环境"],
                    ["渠道费用", "官方平台费用、短信/语音/线路等按实际平台规则计费", "由客户主体承担，或按合同另算"],
                    ["维护人力", "我方持续运维，月费覆盖基础响应和例行复盘", "年度维护费覆盖诊断、升级、远程支持和复盘"],
                    ["升级成本", "我方统一升级更省力", "每个客户环境要单独窗口、备份、验证，成本更高"],
                ],
                [1700, 3830, 3830],
            ),
        ),
        ("h1", "六、推荐落地顺序"),
        ("numbered", [
            "第一批客户优先走托管云端版，降低环境差异和维护成本。",
            "涉及内控或数据本地化要求的客户，再进入私有化部署评估。",
            "私有化客户先做 Lite 试点部署，不直接承诺全渠道、高并发和无人值守。",
            "任何客户正式外发前，都必须完成官方账号授权、白名单测试、人审流程、outbox 门禁、发送回执和异常复盘。",
        ]),
        ("h1", "七、参考依据"),
        ("bullets", [
            "本工程现有文档：P3-05B 托管云端版部署 Runbook、P3-05B 私有化部署包清单、远程维护授权 SOP、P3-05E 企业微信官方 Sandbox Connector。",
            "企业微信开发者中心：https://developer.work.weixin.qq.com/",
            "企业微信微信客服接收消息和事件：https://developer.work.weixin.qq.com/document/path/94670",
            "企业微信微信客服发送消息：https://developer.work.weixin.qq.com/document/path/94677",
        ]),
    ]


def wecom_blocks() -> list[tuple[str, object]]:
    return [
        ("h1", "一、先把概念讲清楚"),
        ("p", "这一步不是注册个人微信机器人，也不是用外挂托管微信号。我们要走的是企业微信/微信客服官方路径：先创建一个微信客服账号，再把一个可调用接口的自建应用绑定到微信客服，最后把企业微信平台的消息回调接到我们的客服中台。"),
        (
            "table",
            (
                ["你看到的名称", "实际含义", "本阶段怎么处理"],
                [
                    ["微信客服账号", "客户在微信里发起咨询时看到的客服入口", "必须创建，并绑定接待人员"],
                    ["接待人员", "企业微信成员，负责人工接待或兜底", "先绑定你自己的企业微信成员"],
                    ["客服链接/二维码", "客户发起咨询的入口", "保存后获取，用个人微信测试"],
                    ["可调用接口的应用", "允许某个自建应用调用微信客服 API", "必须设置，选择测试自建应用"],
                    ["接收消息/回调配置", "企业微信把事件推送到我们的公网 URL", "填 URL、Token、EncodingAESKey"],
                    ["智能机器人", "企业微信官方也有机器人相关能力", "本阶段不依赖它；我们的 AI 在自研客服中台里完成"],
                ],
                [1900, 3500, 3960],
            ),
        ),
        ("callout", ("当前建议", "你现在先不要找“注册机器人”按钮。先完成微信客服账号、接待人员、客服链接、可调用接口应用、回调 URL 验证。机器人式自动回复在我们系统里完成，不靠个人号或群控。", WARN_FILL)),
        ("h1", "二、当前已准备好的我方回调信息"),
        ("p", "当前本地测试使用的是临时公网隧道，适合这次 URL 验证和入站 sandbox。隧道关闭、终端停止或地址变化后，需要重新生成 URL。"),
        (
            "table",
            (
                ["字段", "当前值或来源", "说明"],
                [
                    ["URL", "https://tasty-otters-start.loca.lt/api/webhooks/wecom/channels/1", "填到企业微信后台的回调 URL；只适合本次本地测试"],
                    ["Token", "见本机临时 env 文件里的 WECOM_KF_CALLBACK_TOKEN", "敏感值，不写进聊天和文档正文"],
                    ["EncodingAESKey", "见本机临时 env 文件里的 WECOM_KF_ENCODING_AES_KEY", "敏感值，不写进聊天和文档正文"],
                    ["外发状态", "关闭", "本阶段只接收消息、生成内部会话和 AI 草稿，不真实发回微信"],
                ],
                [1800, 4540, 3020],
            ),
        ),
        ("h1", "三、后台操作总流程"),
        ("numbered", [
            "进入企业微信管理后台，打开“应用管理”。",
            "进入“微信客服”。如果还没启用，先启用微信客服能力。",
            "创建或编辑客服账号，例如“万法常世 AI 客服测试”。",
            "设置头像和名称，保存客服账号。",
            "绑定接待人员，先选择你自己的企业微信成员。",
            "保存后回到客服账号列表，获取客服链接或二维码。",
            "回到微信客服应用详情，找到“可调用接口的应用”，点击“设置”。",
            "选择已有自建应用，或新建一个测试自建应用。",
            "进入该自建应用的详情页，找到接收消息、回调配置、API 接收消息或类似入口。",
            "填入 URL、Token、EncodingAESKey，提交保存，让企业微信验证 URL。",
            "验证通过后，用个人微信打开客服链接，发一条测试消息。",
            "回到我们的客服中台或后端日志，确认消息已进入系统。",
        ]),
        ("h1", "四、详细操作步骤"),
        ("h2", "步骤 1：创建或保存微信客服账号"),
        ("p", "如果你当前页面类似“应用管理 -> 微信客服 -> 应用详情”，并且下面能看到“客服账号”“创建账号”或某个客服账号行，这个页面是对的。先把客服账号本身保存好。"),
        ("numbered", [
            "在“客服账号”区域点击“创建账号”，或点击已有客服账号进入编辑。",
            "客服名称填“万法常世 AI 客服测试”或客户项目名。",
            "头像可以先用默认头像，正式试点再换公司标识。",
            "可见范围先保留当前企业或测试部门。",
            "保存客服账号。",
        ]),
        ("h2", "步骤 2：绑定接待人员"),
        ("p", "接待人员通常在客服账号编辑页或客服账号列表的详情里。不同后台版本按钮可能叫“接待人员”“添加接待人员”“修改接待人员”“接待设置”。"),
        ("numbered", [
            "进入刚创建的客服账号详情。",
            "找到“接待人员”区域。",
            "点击“添加”或“修改”。",
            "选择你自己的企业微信成员。",
            "保存。",
        ]),
        ("callout", ("找不到接待人员时", "先确认你点进的是具体客服账号，而不是微信客服应用总览页；如果页面底部有客服账号列表，点击账号行右侧箭头进入账号详情。", BLUE_FILL)),
        ("h2", "步骤 3：获取客服链接或二维码"),
        ("numbered", [
            "回到客服账号列表。",
            "点击该客服账号行，或点击右侧“接入场景”“客服链接”“二维码”“在微信内接入”等按钮。",
            "复制客服链接或下载二维码。",
            "先不要发给真实客户，只用自己的个人微信测试。",
        ]),
        ("h2", "步骤 4：设置可调用接口的应用"),
        ("p", "企业微信开发者文档明确提到，自建应用需要配置到“微信客服-可调用接口的应用”中，才能调用微信客服相关能力。你截图里已经在“微信客服”应用详情页看到“可调用接口的应用：设置”，这里就是关键入口。"),
        ("numbered", [
            "回到“应用管理 -> 微信客服”的应用详情页。",
            "找到“可调用接口的应用”。",
            "点击“设置”。",
            "选择一个测试自建应用；如果没有，就创建一个新的自建应用，名称可用“万法常世 AI 客服测试应用”。",
            "保存后，记住这个自建应用，因为下一步 URL/Token/AESKey 通常在这个应用的接收消息配置里设置。",
        ]),
        ("h2", "步骤 5：进入自建应用接收消息配置"),
        ("p", "不同企业微信后台版本入口名称不完全一致，你可以按下面顺序找。"),
        ("bullets", [
            "路径 A：应用管理 -> 自建应用 -> 选择刚才绑定的应用 -> 接收消息 -> 设置 API 接收。",
            "路径 B：应用管理 -> 微信客服 -> 可调用接口的应用 -> 设置 -> 进入应用详情 -> 接收消息。",
            "路径 C：应用管理 -> 自建应用 -> 开发者接口 / API / 回调配置 / 接收消息。",
        ]),
        ("h2", "步骤 6：填写 URL、Token、EncodingAESKey"),
        (
            "table",
            (
                ["字段", "怎么填", "注意事项"],
                [
                    ["URL", "填我方给你的公网 HTTPS 回调地址", "必须是 HTTPS；不能填 127.0.0.1；本轮填 localtunnel 地址"],
                    ["Token", "填本机临时 env 文件中的 WECOM_KF_CALLBACK_TOKEN", "要和后端环境变量完全一致"],
                    ["EncodingAESKey", "填本机临时 env 文件中的 WECOM_KF_ENCODING_AES_KEY", "长度通常为 43 位；要和后端完全一致"],
                ],
                [1700, 3900, 3760],
            ),
        ),
        ("callout", ("加解密方式", "如果后台提供消息加解密方式选择，优先使用安全模式或兼容安全模式。本系统已经支持 AES 解密，不要把后台模式改成与后端实现不一致的模式。", BLUE_FILL)),
        ("h2", "步骤 7：提交验证"),
        ("numbered", [
            "确认本地后端还在运行。",
            "确认公网隧道还在运行。",
            "在企业微信后台点击“保存”“提交”或“验证”。",
            "如果验证通过，说明企业微信已经能访问我们的 URL，并且后端能解密返回 echostr。",
            "如果验证失败，先不要反复乱改字段，按本文第六部分排查。",
        ]),
        ("h2", "步骤 8：发送第一条真实测试消息"),
        ("numbered", [
            "用你的个人微信打开客服链接或扫描二维码。",
            "发送一条简单测试消息，例如“你好，我想了解试点部署”。",
            "观察企业微信后台客服会话是否出现。",
            "观察我方后端是否收到消息回调。",
            "确认系统内是否创建会话和消息。",
            "本阶段只生成内部草稿和人工审核，不自动发回微信。",
        ]),
        ("page_break", None),
        ("h1", "五、常见问题和排查"),
        (
            "table",
            (
                ["问题", "可能原因", "处理方式"],
                [
                    ["找不到 URL/Token/AESKey", "还停留在微信客服应用详情，未进入自建应用接收消息配置", "先点“可调用接口的应用：设置”，进入绑定的自建应用详情"],
                    ["提示没有可信 IP", "可信 IP 通常影响自建应用调用 API 的出站能力", "本阶段入站 URL 验证可先走公网隧道；若后台强制要求可信 IP，用云服务器固定 IP"],
                    ["URL 验证失败", "隧道断开、后端停止、Token/AESKey 不一致、URL 路径错误", "先访问 /health，再检查三项字段是否和后端一致"],
                    ["个人微信发消息没进系统", "客服账号未保存、未绑定接待人员、链接不是当前账号、回调配置未生效", "重新确认客服账号、接待人员、客服链接和回调验证状态"],
                    ["担心封号", "使用个人号外挂、Hook、群控、模拟点击会有风险", "只走企业微信/微信客服官方 API，不做个人号托管"],
                ],
                [1900, 3300, 4160],
            ),
        ),
        ("h1", "六、本阶段验收标准"),
        ("bullets", [
            "企业微信后台 URL 验证通过。",
            "个人微信打开客服链接后能发送测试消息。",
            "我方系统收到企业微信加密 XML 消息。",
            "后端验签、AES 解密、XML 解析成功。",
            "系统生成内部会话和消息记录。",
            "后续 AI 草稿进入人工审核。",
            "真实外发仍保持关闭。",
        ]),
        ("h1", "七、官方参考链接"),
        ("bullets", [
            "企业微信开发者中心：https://developer.work.weixin.qq.com/",
            "微信客服接收消息和事件：https://developer.work.weixin.qq.com/document/path/94670",
            "微信客服发送消息：https://developer.work.weixin.qq.com/document/path/94677",
            "获取客服账号链接：https://developer.work.weixin.qq.com/document/path/94692",
            "接待人员管理：https://developer.work.weixin.qq.com/document/path/94693",
            "自建应用与第三方应用的对接：https://developer.work.weixin.qq.com/document/path/95884",
            "自建应用与智能机器人的对接：https://developer.work.weixin.qq.com/document/path/101521",
            "企业微信 API 镜像页，回调配置字段说明：https://qiyeweixin.apifox.cn/doc-417771",
        ]),
    ]


def convert_blocks_for_docx(blocks: list[tuple[str, object]]) -> list[tuple[str, object]]:
    converted: list[tuple[str, object]] = []
    for kind, payload in blocks:
        if kind == "table":
            headers, rows, widths = payload
            converted.append((kind, (headers, rows, widths)))
        elif kind == "callout":
            converted.append((kind, payload))
        else:
            converted.append((kind, payload))
    return converted


def convert_blocks_for_md(blocks: list[tuple[str, object]]) -> list[tuple[str, object]]:
    converted: list[tuple[str, object]] = []
    for kind, payload in blocks:
        if kind == "table":
            headers, rows, _widths = payload
            converted.append((kind, (headers, rows)))
        elif kind == "callout":
            label, text, _fill = payload
            converted.append((kind, (label, text)))
        elif kind == "page_break":
            converted.append((kind, payload))
        else:
            converted.append((kind, payload))
    return converted


def main() -> None:
    docs = [
        (
            DocMeta(
                title="智能客服系统正式部署后运营模式手册",
                subtitle="托管云端版与私有化部署版的账号、运维、维护、成本和责任边界",
                filename="万法常世AI智能客服系统_正式部署后运营模式手册.docx",
                source_filename="万法常世AI智能客服系统_正式部署后运营模式手册.md",
                doc_type="交付与运营手册",
                audience="售前、交付、客户成功、运维负责人、客户项目负责人",
                version="v1.0",
                updated="2026-06-30",
                lead="正式商用只能走托管云端版或私有化部署版；本地公网隧道只用于 sandbox 联调，不能作为客户长期入口。",
            ),
            ops_blocks(),
        ),
        (
            DocMeta(
                title="企业微信微信客服官方接入操作手册",
                subtitle="从创建客服账号、绑定接待人员到配置 URL / Token / EncodingAESKey 的后台操作步骤",
                filename="企业微信微信客服_官方接入操作手册.docx",
                source_filename="企业微信微信客服_官方接入操作手册.md",
                doc_type="渠道接入操作手册",
                audience="项目负责人、企业微信管理员、交付工程师、测试人员",
                version="v1.0",
                updated="2026-06-30",
                lead="不要注册个人号机器人；本阶段按企业微信官方微信客服路径，创建客服账号、绑定接待人员、配置可调用接口应用和回调 URL。",
            ),
            wecom_blocks(),
        ),
    ]
    for meta, blocks in docs:
        md_path = write_markdown(meta, convert_blocks_for_md(blocks))
        docx_path = build_docx(meta, convert_blocks_for_docx(blocks))
        print(f"markdown={md_path}")
        print(f"docx={docx_path}")


if __name__ == "__main__":
    main()
