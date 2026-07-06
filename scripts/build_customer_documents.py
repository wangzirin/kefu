#!/usr/bin/env python3
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output" / "documents"

FONT_LATIN = "Calibri"
FONT_CN = "PingFang SC"
INK = RGBColor(22, 39, 46)
MUTED = RGBColor(92, 105, 111)
ACCENT = RGBColor(21, 113, 127)
ACCENT_DARK = RGBColor(12, 67, 76)
LIGHT_FILL = "F3F7F8"
HEADER_FILL = "E6F0F2"
BORDER = "C9D6DA"
WHITE = "FFFFFF"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


@dataclass
class DocumentMeta:
    title: str
    subtitle: str
    filename: str
    metadata: list[tuple[str, str]]
    lead: str


def set_run_font(
    run,
    *,
    size: float | None = None,
    bold: bool | None = None,
    color: RGBColor | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = FONT_LATIN
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, *, before=0, after=6, line=1.25) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_text(cell, text: str, *, bold=False, color=INK, size=10.5) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=0, line=1.18)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {
        "top": top,
        "bottom": bottom,
        "start": start,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), BORDER)

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            set_cell_width(cell, width)
            set_cell_margins(cell)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[int],
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell_shading(cell, HEADER_FILL)
        set_cell_text(cell, header, bold=True, color=ACCENT_DARK, size=10.3)
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            if idx == 0:
                cell_shading(cell, LIGHT_FILL)
                set_cell_text(cell, value, bold=True, color=INK, size=10.2)
            else:
                cell_shading(cell, WHITE)
                set_cell_text(cell, value, color=INK, size=10.2)
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, after=4)


def add_label_detail_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    add_table(
        doc,
        ["项目", "说明"],
        [[label, detail] for label, detail in rows],
        [2700, CONTENT_WIDTH_DXA - 2700],
    )


def add_section(doc: Document, title: str, body: Iterable[str]) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.add_run(title)
    for text in body:
        add_body(doc, text)


def add_subsection(doc: Document, title: str, body: Iterable[str]) -> None:
    p = doc.add_paragraph(style="Heading 2")
    p.add_run(title)
    for text in body:
        add_body(doc, text)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=7, line=1.25)
    run = p.add_run(text)
    set_run_font(run, size=11.5, color=INK)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph_spacing(p, after=4, line=1.25)
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        run = p.add_run(item)
        set_run_font(run, size=11.2, color=INK)


def add_callout(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.rows[0].cells[0]
    cell_shading(cell, LIGHT_FILL)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    cell.text = ""
    title_p = cell.paragraphs[0]
    set_paragraph_spacing(title_p, after=4, line=1.15)
    title_run = title_p.add_run(title)
    set_run_font(title_run, size=11.5, bold=True, color=ACCENT_DARK)
    body_p = cell.add_paragraph()
    set_paragraph_spacing(body_p, after=0, line=1.22)
    body_run = body_p.add_run(text)
    set_run_font(body_run, size=10.8, color=INK)
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, after=4)


def configure_doc(doc: Document, running_label: str) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(11.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, ACCENT_DARK, 18, 10),
        ("Heading 2", 13, ACCENT, 14, 7),
        ("Heading 3", 12, ACCENT_DARK, 10, 5),
    ]:
        style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    list_style = styles["List Bullet"]
    list_style.font.name = FONT_LATIN
    list_style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    list_style.font.size = Pt(11.2)
    list_style.font.color.rgb = INK
    list_style.paragraph_format.space_after = Pt(4)
    list_style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run(running_label)
    set_run_font(run, size=9.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("万法常世 AI 全智能客服系统")
    set_run_font(footer_run, size=9.5, color=MUTED)


def add_cover(doc: Document, meta: DocumentMeta) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=2, line=1.0)
    run = p.add_run("万法常世")
    set_run_font(run, size=12, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=7, line=1.06)
    run = p.add_run(meta.title)
    set_run_font(run, size=26, bold=True, color=ACCENT_DARK)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=18, line=1.2)
    run = p.add_run(meta.subtitle)
    set_run_font(run, size=13.5, color=MUTED)

    left = meta.metadata[: (len(meta.metadata) + 1) // 2]
    right = meta.metadata[(len(meta.metadata) + 1) // 2 :]
    table = doc.add_table(rows=max(len(left), len(right)), cols=4)
    set_table_geometry(table, [1600, 3080, 1600, 3080])
    for row_idx in range(len(table.rows)):
        values = []
        if row_idx < len(left):
            values.extend(left[row_idx])
        else:
            values.extend(["", ""])
        if row_idx < len(right):
            values.extend(right[row_idx])
        else:
            values.extend(["", ""])
        for cell_idx, value in enumerate(values):
            cell = table.rows[row_idx].cells[cell_idx]
            cell_shading(cell, LIGHT_FILL if cell_idx in {0, 2} else WHITE)
            set_cell_text(
                cell,
                value,
                bold=cell_idx in {0, 2},
                color=ACCENT_DARK if cell_idx in {0, 2} else INK,
                size=9.8,
            )

    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, after=10)
    add_callout(doc, "文件说明", meta.lead)


def build_product_intro() -> None:
    meta = DocumentMeta(
        title="AI 全智能客服系统产品介绍",
        subtitle="面向企业接待、知识库问答、人工审核和质量运营的一体化客服系统",
        filename="万法常世AI全智能客服系统_产品介绍_正式版.docx",
        metadata=[
            ("适用对象", "企业负责人、客服负责人、运营负责人、数字化负责人"),
            ("适用版本", "Lite 试点版、标准运营版"),
            ("交付形态", "私有化部署、受控试点、后续可扩展渠道"),
            ("文件日期", "2026年6月29日"),
        ],
        lead="这份材料用于介绍系统的业务价值、使用场景、核心能力、版本差异和上线边界。文档面向采购评估和项目启动沟通，不包含内部报价策略或客户环境密钥。",
    )
    doc = Document()
    configure_doc(doc, "AI 全智能客服系统 | 产品介绍")
    add_cover(doc, meta)

    add_section(
        doc,
        "产品定位",
        [
            "万法常世 AI 全智能客服系统是一套面向企业客服和运营团队的知识库型客服系统。系统把企业资料、业务规则、常见问题、人工审核和质量复盘放在同一个工作台中，让接待过程更一致、更容易追踪，也更适合长期维护。",
            "系统的重点不是用机器人替代所有人工，而是让常见问题快速生成有依据的回复草稿，把复杂、高风险、证据不足的问题交给人工处理，并把每一次审核、改写和发送尝试记录下来。",
        ],
    )
    add_label_detail_table(
        doc,
        [
            ("适合解决的问题", "重复咨询多、客服口径不统一、知识分散、人工培训成本高、售后和合同问题需要留痕。"),
            ("系统处理方式", "先检索知识库，再生成回复草稿；草稿带引用、风险提示和人工审核入口。"),
            ("上线原则", "先试点一个入口和一批高频问题，再逐步扩展渠道、知识库和坐席规模。"),
        ],
    )

    add_section(
        doc,
        "主要使用场景",
        [
            "系统适合从一个明确入口开始，例如官网咨询、企业微信客服、公众号咨询或指定私域入口。试点稳定后，再根据官方接口和平台授权逐步扩展到更多渠道。",
        ],
    )
    add_table(
        doc,
        ["场景", "适用情况", "系统价值"],
        [
            ["官网咨询", "访客询问产品、价格、合作流程和售后政策", "先接住咨询，再把高意向问题和知识缺口沉淀下来"],
            ["私域接待", "企业微信、社群或顾问团队需要统一口径", "减少重复解释，保持价格、活动和交付口径一致"],
            ["售前转化", "需要快速回答套餐、案例、交付周期和合作条件", "草稿带知识来源，坐席可以更快确认并回复"],
            ["售后服务", "退款、保修、发票、合同、账号权限等问题较多", "高风险问题自动提示人工审核，降低误承诺风险"],
            ["多坐席协作", "多人共同处理咨询，需要分工和追踪", "工作台集中展示待审核、待发送、失败和高风险会话"],
            ["知识运营", "资料更新频繁，容易出现旧口径", "通过题库回归和知识缺口复盘发现需要更新的内容"],
        ],
        [1700, 3400, 4260],
    )

    add_section(
        doc,
        "系统如何工作",
        [
            "一次咨询进入系统后，会先被识别为可信入站消息，再进入知识检索和回复草稿生成流程。坐席在工作台查看草稿和引用证据，确认后才进入待发送区。系统保留审核和发送记录，便于后续复盘。",
        ],
    )
    add_table(
        doc,
        ["步骤", "系统动作", "使用团队看到什么"],
        [
            ["接收咨询", "官网或授权渠道把咨询送入系统，系统记录来源和会话", "新会话、渠道来源、联系人信息"],
            ["检索知识", "从知识卡片和文档片段中查找相关资料", "命中的知识来源、引用片段、置信度"],
            ["生成草稿", "根据问题和证据生成可审核回复", "AI 草稿、风险等级、转人工原因"],
            ["人工审核", "坐席批准、改写、拒绝或转交业务负责人", "待审核队列、审计记录、改写历史"],
            ["发送确认", "通过外发门禁后进入发送流程", "待发送、发送失败、回执和失败原因"],
            ["质量复盘", "定期检查题库回归、知识缺口和坐席改写", "质量摘要、问题分类、知识更新建议"],
        ],
        [1400, 3900, 4060],
    )

    add_section(
        doc,
        "核心能力",
        [
            "系统以知识库、AI 草稿、人工审核、外发门禁和质量复盘为核心。每个能力都围绕实际接待流程设计，避免只停留在固定问答或单次聊天。",
        ],
    )
    add_table(
        doc,
        ["能力", "说明", "适用价值"],
        [
            ["自建知识库", "支持知识卡片、文档导入、分块检索、引用来源和版本管理", "把企业内部资料变成可检索、可复盘的接待依据"],
            ["AI 回复草稿", "根据用户问题和知识证据生成回复建议", "减少重复打字，提高常见问题处理速度"],
            ["引用证据", "草稿展示使用了哪些知识片段", "坐席可以判断回复是否有依据"],
            ["人工审核", "复杂、高风险、低置信问题进入人工处理", "避免系统直接承诺退款、赔付、合同或法律判断"],
            ["待发送门禁", "审核通过和真实发送分开处理", "降低误发、错发和渠道权限风险"],
            ["质量复盘", "通过题库、线上反馈、改写率和缺口记录看质量变化", "持续发现知识过期和新问题"],
            ["模型路由", "简单问题低成本处理，普通问题走标准模型，高风险问题保留人工", "控制成本，也保留安全边界"],
            ["诊断与运维", "提供健康检查、日志、备份、恢复、诊断包和故障响应方案", "上线后可以维护、排障和追踪责任"],
        ],
        [1750, 4300, 3310],
    )

    add_section(
        doc,
        "版本选择",
        [
            "建议企业先按业务风险和团队规模选择版本。Lite 试点版适合快速验证效果，标准运营版适合进入日常运营。",
        ],
    )
    add_table(
        doc,
        ["版本", "适合企业", "包含内容", "交付重点"],
        [
            ["Lite 试点版", "单业务线、小团队、希望先验证智能客服价值", "一个入口、轻量知识库、AI 草稿、人工审核、基础题库、基础部署和诊断", "快速上线试点，先验证常见问题、知识命中和坐席使用意愿"],
            ["标准运营版", "多坐席团队、需要长期维护知识库和质量指标", "账号权限、工作台、知识运营、质量中心、模型路由、渠道沙盒、审计、备份恢复", "进入正式运营，支持持续维护和扩展"],
        ],
        [1600, 2500, 3300, 1960],
    )

    add_section(
        doc,
        "安全和合规边界",
        [
            "客服系统直接接触企业用户、价格、售后和平台规则，因此默认采用谨慎上线方式。系统不会默认绕过人工审核，也不会使用个人号外挂、群控或模拟点击作为正式接入方式。",
        ],
    )
    add_table(
        doc,
        ["边界", "系统默认做法"],
        [
            ["真实外发", "默认关闭。上线初期先做草稿和人工审核，真实发送需要双方确认。"],
            ["高风险问题", "退款、赔付、合同、投诉、法务和隐私问题默认进入人工处理。"],
            ["客户数据", "真实数据进入系统前需要授权和脱敏，文档和诊断包不保存密钥或聊天原文。"],
            ["渠道接入", "只走官网自有入口、官方 API、开放平台或服务商授权。"],
            ["准确率", "不承诺 100% 准确，也不建议无人监管；通过题库和人工反馈持续维护。"],
        ],
        [2200, CONTENT_WIDTH_DXA - 2200],
    )

    add_section(
        doc,
        "推荐上线节奏",
        [
            "更稳妥的上线方式是先试点，再扩展。前两周建议保留人工审核，重点观察常见问题命中、坐席改写、高风险转人工和知识缺口。",
        ],
    )
    add_table(
        doc,
        ["阶段", "时间", "重点工作", "通过标准"],
        [
            ["准备", "第 1 周", "确认入口、知识资料、角色、禁用话术和试点目标", "资料完整，风险边界清楚"],
            ["部署", "第 1 周", "部署系统、配置账号、导入知识库、设置外发开关", "系统可登录，健康检查通过"],
            ["试运营", "第 2-3 周", "所有 AI 草稿先人工审核，收集知识缺口和改写记录", "常见问题可处理，高风险能转人工"],
            ["复盘", "第 4 周", "查看质量指标、修订知识库、决定是否扩大范围", "明确下一批渠道、坐席或知识更新计划"],
        ],
        [1400, 1400, 4200, 2360],
    )

    add_callout(
        doc,
        "采购建议",
        "如果企业第一次使用知识库型智能客服，建议从 Lite 试点版开始，先验证一个入口、一批高频问题和一个坐席流程。试点稳定后，再升级标准运营版，扩展权限、质量中心、更多渠道和持续运维服务。",
    )
    doc.save(OUTPUT_DIR / meta.filename)


def build_service_system() -> None:
    meta = DocumentMeta(
        title="服务体系介绍",
        subtitle="从试点部署到知识运营、故障响应和质量维护的服务安排",
        filename="万法常世AI全智能客服系统_服务体系介绍_正式版.docx",
        metadata=[
            ("适用对象", "项目负责人、客服负责人、运营负责人、技术负责人"),
            ("适用版本", "Lite 试点版、标准运营版"),
            ("服务范围", "实施、培训、运维、质量复盘、渠道支持"),
            ("文件日期", "2026年6月29日"),
        ],
        lead="这份材料说明系统交付后的服务方式、双方职责、维护节奏和故障处理口径。文档适合项目启动、采购评估和售后服务确认使用。",
    )
    doc = Document()
    configure_doc(doc, "AI 全智能客服系统 | 服务体系")
    add_cover(doc, meta)

    add_section(
        doc,
        "服务目标",
        [
            "智能客服上线后的效果，取决于知识资料、业务口径、渠道权限、模型配置和日常复盘。服务体系的目标，是让系统不仅能安装，还能被使用团队稳定接手、持续维护，并在出现问题时快速定位。",
            "服务过程会把部署、知识整理、坐席培训、质量检查、故障处理和月度复盘放在同一套节奏里，避免系统上线后无人维护、知识过期或责任不清。",
        ],
    )
    add_label_detail_table(
        doc,
        [
            ("实施目标", "完成试点环境部署、账号配置、知识导入、入口接入和上线前检查。"),
            ("运营目标", "让使用团队掌握草稿审核、知识更新、质量复盘和高风险处理。"),
            ("维护目标", "保证系统有健康检查、日志、备份、恢复、诊断包和故障响应方式。"),
        ],
    )

    add_section(
        doc,
        "服务内容",
        [
            "服务内容按项目推进顺序组织。Lite 试点版更重视快速上线和低风险验证；标准运营版更重视权限、质量、渠道、审计和持续维护。",
        ],
    )
    add_table(
        doc,
        ["模块", "服务内容", "交付结果"],
        [
            ["需求确认", "确认业务范围、试点入口、使用人员、知识资料、风险边界", "项目启动清单、试点范围、资料清单"],
            ["系统部署", "部署前后端、数据库、Redis、模型配置、基础账号和访问域名", "部署记录、环境变量清单、健康检查结果"],
            ["知识库建设", "整理资料，导入知识卡片和文档，标记禁用话术和风险边界", "知识库初版、来源清单、版本记录"],
            ["题库回归", "准备 50-100 条脱敏问题，标注期望证据和人审边界", "回归结果、知识缺口、人工复核建议"],
            ["坐席培训", "讲解待办队列、草稿审核、改写、转人工和待发送确认", "培训记录、操作确认、常见问题处理口径"],
            ["试运营陪跑", "观察真实咨询、收集改写、修订知识库、处理低置信问题", "周复盘、知识更新清单、风险问题记录"],
            ["运维维护", "健康检查、日志、备份恢复、升级回滚、诊断包和故障响应", "维护记录、故障处理单、恢复演练记录"],
        ],
        [1700, 4300, 3360],
    )

    add_section(
        doc,
        "实施流程",
        [
            "项目建议以四个阶段推进。每个阶段都有明确输入和输出，避免资料、账号、外发权限和质量验收混在一起处理。",
        ],
    )
    add_table(
        doc,
        ["阶段", "重点工作", "需要使用团队配合", "完成标志"],
        [
            ["启动", "明确业务范围、入口、角色、风险问题和试点目标", "提供资料清单、使用人员、业务负责人", "启动清单确认"],
            ["部署", "安装系统、配置数据库、账号、模型和访问域名", "提供服务器或部署环境、域名和安全要求", "健康检查通过"],
            ["知识", "整理知识、导入系统、建立题库、运行回归", "确认价格、售后、合同、禁用话术和生效时间", "知识库初版可用"],
            ["试运营", "坐席审核草稿，复盘低置信和知识缺口", "按日处理待办，按周确认知识修订", "形成第一份试运营复盘"],
        ],
        [1350, 3150, 3150, 1710],
    )

    add_section(
        doc,
        "运维维护",
        [
            "系统交付后，维护重点不只是服务器在线，还包括知识库是否过期、模型是否可用、渠道是否失败、草稿是否被频繁改写。维护项需要固定节奏。",
        ],
    )
    add_table(
        doc,
        ["维护项", "检查内容", "建议频率"],
        [
            ["健康检查", "前端、后端、数据库、Redis、模型 provider、渠道回调", "每日或每周"],
            ["日志与审计", "登录、审核、发送尝试、失败复盘、知识更新", "每周"],
            ["备份恢复", "数据库、知识库、配置、题库和输出报告", "试点每日备份，恢复演练按月"],
            ["成本观察", "模型调用量、失败率、延迟、估算费用", "每周或每月"],
            ["知识质量", "题库回归、引用支撑、改写率、知识缺口", "每周或每次知识更新后"],
            ["渠道状态", "回调、验签、授权、发送失败和平台回执", "上线后高频检查"],
        ],
        [1800, 5200, 2360],
    )

    add_section(
        doc,
        "故障响应",
        [
            "故障处理先止血，再定位，再复盘。涉及真实外发、数据恢复、模型异常和渠道权限时，应优先保护业务安全和客户数据。",
        ],
    )
    add_table(
        doc,
        ["等级", "示例", "处理目标"],
        [
            ["P1 紧急", "系统不可用、错误外发风险、数据恢复、重大安全风险", "立即止血，关闭风险开关，优先恢复核心使用"],
            ["P2 高", "渠道大面积失败、模型不可用、知识库发布错误", "当日定位并给出修复或绕行方案"],
            ["P3 中", "单个功能异常、部分坐席无法使用、个别题库回归下降", "按维护窗口处理，并记录复盘"],
            ["P4 低", "文案、普通配置、使用咨询、培训补充", "排期处理"],
        ],
        [1200, 4300, 3860],
    )
    add_callout(
        doc,
        "故障记录要求",
        "每次故障都应记录发生时间、影响范围、临时处理、根因、修复动作、是否需要知识库更新，以及后续预防措施。涉及客户生产环境的操作必须有授权和回滚点。",
    )

    add_section(
        doc,
        "质量运营",
        [
            "智能客服的质量会随产品、价格、活动、平台规则和用户问法变化而变化。质量运营的重点，是及早发现下降趋势，而不是等投诉出现后再处理。",
        ],
    )
    add_table(
        doc,
        ["指标", "说明", "建议处理方式"],
        [
            ["知识命中率", "问题是否能找到相关知识", "下降时检查知识是否缺失、过期或分块不合理"],
            ["引用支撑率", "草稿是否被知识证据支持", "下降时检查检索和提示词，必要时转人工"],
            ["转人工率", "多少问题进入人工处理", "异常升高时检查知识缺口和模型/provider 状态"],
            ["坐席改写率", "草稿被人工大幅修改的比例", "升高时复查口径、语气和证据完整度"],
            ["知识缺口", "系统无法可靠回答的问题", "按高频和高风险优先补齐"],
            ["发送失败率", "渠道或权限导致的发送异常", "检查官方授权、回执和外发开关"],
        ],
        [1800, 3800, 3760],
    )

    add_section(
        doc,
        "双方职责",
        [
            "服务方负责系统、部署、维护和质量支持；使用团队负责业务资料、政策口径、授权和高风险审核。两个角色配合得越清楚，系统越容易稳定。",
        ],
    )
    add_table(
        doc,
        ["责任方", "主要职责"],
        [
            ["服务方", "系统部署、技术维护、知识导入支持、题库回归、质量复盘、故障排查、升级回滚方案。"],
            ["使用团队", "提供真实业务资料，确认产品、价格、售后、合同和禁用话术，审核高风险回复，提供脱敏题库，授权渠道和模型调用。"],
            ["双方共同确认", "试点目标、上线范围、外发开关、自动回复开放边界、月度复盘和后续扩展计划。"],
        ],
        [2100, CONTENT_WIDTH_DXA - 2100],
    )

    add_section(
        doc,
        "服务分层",
        [
            "不同企业需要的服务强度不同。建议把一次性交付和持续维护分开确认，避免系统上线后缺少知识更新、质量复盘和故障响应预算。",
        ],
    )
    add_table(
        doc,
        ["服务层级", "适合版本", "服务内容"],
        [
            ["基础维护", "Lite 试点版", "健康检查、基础故障处理、月度质量摘要、少量知识小修和诊断包支持"],
            ["标准运营", "标准运营版", "周/月质量复盘、知识更新、渠道排障、模型策略复核、备份恢复和上线陪跑"],
            ["专项支持", "标准运营版及以上", "私有化部署、多渠道官方接入、高并发、合规审计、专属响应和专项优化"],
        ],
        [1800, 2200, 5360],
    )

    add_callout(
        doc,
        "上线建议",
        "试点前两周建议保持所有 AI 草稿人工审核。低风险自动回复应在题库回归、坐席抽查、业务负责人确认和外发审批通过后再逐步开放。",
    )
    doc.save(OUTPUT_DIR / meta.filename)


def build_user_manual() -> None:
    meta = DocumentMeta(
        title="使用手册",
        subtitle="日常接待、草稿审核、知识维护和质量复盘操作说明",
        filename="万法常世AI全智能客服系统_使用手册_正式版.docx",
        metadata=[
            ("适用对象", "管理员、运营负责人、坐席、业务负责人、技术维护人员"),
            ("适用版本", "Lite 试点版、标准运营版"),
            ("使用场景", "试点上线、坐席培训、日常操作、质量复盘"),
            ("文件日期", "2026年6月29日"),
        ],
        lead="这份手册用于帮助使用团队完成系统登录、待办处理、AI 草稿审核、知识库维护、质量复盘和上线前检查。手册不包含真实账号、密钥或客户数据。",
    )
    doc = Document()
    configure_doc(doc, "AI 全智能客服系统 | 使用手册")
    add_cover(doc, meta)

    add_section(
        doc,
        "系统能做什么",
        [
            "系统用于把咨询入口、企业知识、AI 回复草稿、人工审核和发送记录放进同一个工作台。使用团队可以在工作台处理待审核草稿、待发送回复、失败任务、高风险会话和知识缺口。",
            "系统不是无人监管的自动回复工具。试点期建议所有草稿先经过人工审核，确认知识库和渠道外发规则稳定后，再讨论是否开放低风险自动回复。",
        ],
    )
    add_label_detail_table(
        doc,
        [
            ("接收咨询", "接收官网或授权渠道的入站消息，并记录来源。"),
            ("生成草稿", "根据知识库证据生成回复建议，展示引用来源和风险提示。"),
            ("人工审核", "坐席确认、改写、拒绝或转人工处理。"),
            ("发送确认", "审核通过后进入待发送区，真实外发受开关和渠道授权控制。"),
            ("质量复盘", "通过题库回归、改写率、知识缺口和失败记录判断质量变化。"),
        ],
    )

    add_section(
        doc,
        "角色说明",
        [
            "不同角色看到的工作重点不同。项目启动时应先确认账号、权限和负责人，避免知识、渠道和高风险问题无人确认。",
        ],
    )
    add_table(
        doc,
        ["角色", "主要职责"],
        [
            ["管理员", "配置账号、权限、渠道、模型、知识库和运维设置。"],
            ["运营负责人", "维护知识库，查看质量数据，组织低置信和知识缺口复盘。"],
            ["坐席", "审核 AI 草稿，改写回复，处理高风险和转人工会话。"],
            ["业务负责人", "确认产品政策、价格、售后、合同、发票和禁用话术。"],
            ["技术维护人员", "处理部署、备份、恢复、升级、日志和故障排查。"],
        ],
        [2100, CONTENT_WIDTH_DXA - 2100],
    )

    add_section(
        doc,
        "每日工作流程",
        [
            "日常工作建议从待办队列开始。系统会把需要处理的问题集中在工作台，而不是让坐席逐条翻所有历史消息。",
        ],
    )
    add_table(
        doc,
        ["队列", "处理方式"],
        [
            ["待人工审核", "查看用户原始问题、AI 草稿、引用证据和风险等级；可以批准、改写或拒绝。"],
            ["待发送确认", "发送前确认联系人、渠道、内容和外发开关。"],
            ["发送失败", "查看失败原因，判断是否重试、改写、转人工或联系技术维护。"],
            ["知识缺口", "整理缺失问题，提交业务负责人确认后补充知识库。"],
            ["高风险会话", "退款、赔付、合同、投诉、舆情、隐私等问题必须人工处理。"],
            ["最近评测运行", "知识更新后查看题库回归结果是否下降。"],
            ["渠道健康", "检查官网、企业微信、公众号等入口是否正常。"],
        ],
        [2200, CONTENT_WIDTH_DXA - 2200],
    )

    add_subsection(
        doc,
        "审核 AI 草稿",
        [
            "审核草稿时，不要只看回复文字。必须同时看用户原始问题、引用证据、风险等级、转人工原因、历史会话和审计记录。",
        ],
    )
    add_table(
        doc,
        ["可以批准", "需要改写或转人工"],
        [
            ["低风险常见咨询，知识库证据明确，草稿没有超出政策范围。", "没有引用证据，或证据与草稿不一致。"],
            ["产品、交付、账号、基础售后等问题能被现有知识支撑。", "涉及退款、赔付、合同解释、投诉、法务、隐私或重大承诺。"],
            ["语气符合企业服务口径，不承诺最低价、必赔或无法保证的结果。", "用户问题超出知识库范围，或需要业务负责人确认。"],
        ],
        [4680, 4680],
    )

    add_subsection(
        doc,
        "处理待发送",
        [
            "审核通过不等于已经发送。待发送区用于最后确认回复内容和渠道状态。如果系统显示外发关闭，代表当前只做内部审核和演示，不会向真实平台发送。",
        ],
    )
    add_bullets(
        doc,
        [
            "确认回复内容没有敏感承诺、错误价格或过期政策。",
            "确认收件人、渠道和会话来源正确。",
            "确认该渠道已获授权，且外发开关符合本次试点安排。",
            "发送失败时先看错误原因，不要重复点击造成重复消息。",
        ],
    )

    add_section(
        doc,
        "知识库维护",
        [
            "知识库是系统效果的核心。产品资料、价格、活动、售后规则、合同条款和禁用话术发生变化后，应按版本更新，不建议无记录覆盖旧内容。",
        ],
    )
    add_table(
        doc,
        ["步骤", "操作", "注意事项"],
        [
            ["整理资料", "收集产品、价格、售后、交付、账号、合同和发票口径", "明确负责人、生效时间和适用范围"],
            ["转成知识", "整理为知识卡片或文档", "避免把危险承诺写成可直接引用的答案"],
            ["导入系统", "导入后生成检索片段和引用", "检查是否能按高频问题检索到"],
            ["题库回归", "使用脱敏题库检查命中、引用和转人工", "质量下降时先不要发布"],
            ["人工抽查", "业务负责人抽查高频和高风险问题", "通过后记录版本并发布"],
            ["发布复盘", "观察改写率、缺口和负反馈", "必要时回滚或补充知识"],
        ],
        [1400, 3900, 4060],
    )

    add_section(
        doc,
        "质量复盘",
        [
            "质量复盘不只看系统回答了多少问题，更要看草稿是否有证据、坐席是否频繁改写、高风险是否转人工、知识缺口是否被补齐。",
        ],
    )
    add_table(
        doc,
        ["指标", "含义", "出现异常时先看什么"],
        [
            ["知识命中率", "用户问题能否找到相关知识", "资料是否缺失、过期或分块不合理"],
            ["引用支撑率", "AI 草稿是否被知识证据支持", "检索结果和草稿是否一致"],
            ["转人工率", "多少问题需要人工处理", "风险规则、知识缺口、模型可用性"],
            ["低置信率", "系统不确定的问题比例", "新业务、新活动、新渠道来源"],
            ["坐席改写率", "草稿被大幅修改的比例", "语气、政策、证据和提示词"],
            ["知识缺口数", "知识库没有覆盖的问题数量", "高频缺口优先补齐"],
            ["发送失败数", "渠道或权限导致的发送异常", "授权、回执、外发开关和网络状态"],
        ],
        [1700, 3400, 4260],
    )

    add_section(
        doc,
        "准确率下降时怎么处理",
        [
            "如果发现草稿质量下降、坐席改写变多或投诉增加，建议按固定顺序排查，不要只归因于模型。",
        ],
    )
    add_bullets(
        doc,
        [
            "先检查最近是否更新了产品、价格、活动或售后政策。",
            "检查知识库是否刚发布新版本，是否漏掉旧问题。",
            "检查模型 provider、模型名称或参数是否变更。",
            "检查是否新增渠道，用户问法是否明显变化。",
            "抽查低置信、高改写和高风险会话。",
            "补充知识库或调整禁用话术后，重新跑题库回归。",
            "通过后再发布；未通过时继续保持人工审核。",
        ],
    )

    add_section(
        doc,
        "上线前检查清单",
        [
            "上线前应由项目负责人、业务负责人、运营负责人和技术维护人员共同确认。任何真实外发都必须在授权、测试和风险确认后开启。",
        ],
    )
    add_table(
        doc,
        ["检查项", "负责人", "状态"],
        [
            ["管理员账号确认", "使用团队", "待确认"],
            ["坐席账号确认", "使用团队", "待确认"],
            ["产品资料整理", "业务负责人", "待确认"],
            ["价格和套餐确认", "业务负责人", "待确认"],
            ["售后政策确认", "业务负责人", "待确认"],
            ["禁用话术确认", "业务负责人", "待确认"],
            ["50-100 条脱敏问题", "使用团队 + 服务方", "待确认"],
            ["知识库导入", "服务方", "待执行"],
            ["质量回归", "服务方", "待执行"],
            ["渠道授权", "使用团队", "待确认"],
            ["外发开关", "双方", "默认关闭"],
            ["备份和诊断包", "技术维护人员", "待确认"],
        ],
        [4200, 3100, 2060],
    )

    add_section(
        doc,
        "常见问题",
        [
            "下面问题适合放在培训和试运营启动会上统一说明。",
        ],
    )
    add_label_detail_table(
        doc,
        [
            ("AI 会不会自己乱发消息", "系统默认以人工审核和外发门禁为边界。未经配置和授权，草稿不会直接发送到真实渠道。"),
            ("知识更新后多久生效", "轻量知识通常较快生效；文档知识建议先跑题库回归，再发布给坐席使用。"),
            ("为什么有些问题转人工", "常见原因包括低置信、无知识命中、退款赔付、合同法务、投诉舆情、隐私信息或超出知识库范围。"),
            ("能否接多个平台", "可以逐步接入，但必须走官方 API、开放平台或服务商授权。个人号外挂、群控和模拟点击不属于正式方案。"),
            ("上线后谁维护", "使用团队负责业务口径和高风险确认；服务方负责系统部署、技术维护、质量复盘支持和故障处理。"),
        ],
    )

    add_callout(
        doc,
        "试点期建议",
        "前两周保持所有 AI 草稿人工审核。每天复盘知识缺口，每周看一次质量数据。先稳定一个入口，再扩展更多渠道和坐席。",
    )
    doc.save(OUTPUT_DIR / meta.filename)


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    build_product_intro()
    build_service_system()
    build_user_manual()
    for path in sorted(OUTPUT_DIR.glob("万法常世AI全智能客服系统_*_正式版.docx")):
        print(f"WROTE {path}")


if __name__ == "__main__":
    main()
